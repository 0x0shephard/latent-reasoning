from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("output/documents/Accuracy_Bearing_Eigenspaces_for_Efficient_CODI_Inference_NeurIPS_Style.docx")

BLACK = RGBColor(0, 0, 0)
FONT = "Times New Roman"


def set_cell_border(cell, **edges):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in edges:
            continue
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edges[edge].items():
            element.set(qn("w:" + key), str(value))


def set_cell_margins(cell, top=20, start=70, bottom=20, end=70):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + margin))
        if node is None:
            node = OxmlElement("w:" + margin)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run(run, size=10, bold=False, italic=False, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    run.bold = bold
    run.italic = italic


def add_para(doc, text="", *, size=10, bold=False, italic=False, align=None,
             before=0, after=5.5, line=Pt(11), keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    p.paragraph_format.keep_with_next = keep
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    size = {1: 12, 2: 10, 3: 10}[level]
    before = {1: 12, 2: 9, 3: 7}[level]
    after = {1: 6, 2: 4, 3: 3}[level]
    p = add_para(doc, text, size=size, bold=True, before=before, after=after, keep=True)
    p.style = doc.styles[f"Heading {level}"]
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(11)
    run = p.add_run(text)
    set_run(run)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(11)
    run = p.add_run(text)
    set_run(run)
    return p


def add_formula(doc, text):
    p = add_para(doc, text, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=6)
    return p


def add_horizontal_rule(doc, points):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(points * 8)))
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    return p


def apply_booktabs(table):
    thin = {"val": "single", "sz": "4", "space": "0", "color": "000000"}
    thick = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=thick, bottom=thin)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=thick)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_results_table(doc):
    headers = ["Condition", "Dimensions", "Variance share", "Exact match", "Baseline retained"]
    rows = [
        ["Baseline", "768", "100.00%", "0.4337", "1.000"],
        ["Retain 0-3", "4", "82.31%", "0.0265", "0.061"],
        ["Retain 4-15", "12", "7.49%", "0.2191", "0.505"],
        ["Retain 4-31", "28", "11.31%", "0.3806", "0.878"],
        ["Retain 0-31", "32", "93.62%", "0.4094", "0.944"],
        ["Remove 4-31", "28 removed", "11.31%", "0.1289", "0.297"],
        ["Remove 0-3", "4 removed", "82.31%", "0.4185", "0.965"],
    ]
    add_para(doc, "Table 1: Confirmatory full-GSM8K results for interventions at the answer-cue endpoint.", size=9, after=3)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.80), Inches(0.70), Inches(0.95), Inches(0.90), Inches(1.15)]
    for i, (cell, label) in enumerate(zip(table.rows[0].cells, headers)):
        cell.width = widths[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_run(run, size=8, bold=True)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, (cell, value) in enumerate(zip(cells, row)):
            cell.width = widths[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(value)
            set_run(run, size=8)
    apply_booktabs(table)
    add_para(doc, "", size=8, after=2)


def add_efficiency_table(doc):
    headers = ["Readout", "Online MACs per token", "Theoretical reduction", "T4 time"]
    rows = [
        ["Full 768-dimensional head", "38,597,376", "1.00x", "1,327 microseconds"],
        ["Rank-28 eigenspace head", "1,428,700", "27.0x fewer", "120 microseconds"],
        ["Rank-32 eigenspace head", "1,632,800", "23.6x fewer", "117 microseconds"],
        ["Rank-64 eigenspace head", "3,265,600", "11.8x fewer", "154 microseconds"],
    ]
    add_para(doc, "Table 2: Isolated language-model-head measurements on a Tesla T4 in float32. These are component timings, not end-to-end generation speedups.", size=9, after=3)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.70), Inches(1.35), Inches(1.15), Inches(1.30)]
    for i, (cell, label) in enumerate(zip(table.rows[0].cells, headers)):
        cell.width = widths[i]
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_run(run, size=8, bold=True)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, (cell, value) in enumerate(zip(cells, row)):
            cell.width = widths[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run(run, size=8)
    apply_booktabs(table)
    add_para(doc, "", size=8, after=2)


def add_timeline_table(doc):
    headers = ["Phase", "Work", "Planned output"]
    rows = [
        ["1", "Reproduce all preliminary measurements and freeze data splits.", "Audited baseline and artifact manifest"],
        ["2", "Implement rank-28 and rank-32 heads in the generation loop.", "End-to-end inference implementation"],
        ["3", "Run rank sweeps and comparison methods.", "Accuracy, agreement, and latency results"],
        ["4", "Evaluate transfer across tasks, seeds, and checkpoints.", "Generalization and failure analysis"],
        ["5", "Write the final study and release reproducible artifacts.", "Dissertation paper and code package"],
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(0.50), Inches(3.00), Inches(2.00)]
    for i, (cell, label) in enumerate(zip(table.rows[0].cells, headers)):
        cell.width = widths[i]
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(label), size=8, bold=True)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, (cell, value) in enumerate(zip(cells, row)):
            cell.width = widths[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(value), size=8)
    apply_booktabs(table)


def add_reference(doc, number, text, url):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(10)
    set_run(p.add_run(f"[{number}] {text} "), size=9)
    set_run(p.add_run(url), size=9)


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.5)
    section.header_distance = Pt(25)
    section.footer_distance = Pt(30)
    section.different_first_page_header_footer = True

    for style_name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
        style.font.color.rgb = BLACK
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.line_spacing = Pt(11)

    # NeurIPS 2026 title block: 17-point title between 4-point and 1-point rules.
    add_horizontal_rule(doc, 4)
    add_para(
        doc,
        "Accuracy-Bearing Eigenspaces for Efficient CODI Inference",
        size=17,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=18,
        after=18,
        line=Pt(20),
    )
    add_horizontal_rule(doc, 1)
    add_para(doc, "Muhammad Jon Raza", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=2)
    add_para(doc, "Roll Number: 27100293", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)

    add_para(doc, "Abstract", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=4, keep=True)
    abstract = add_para(doc, (
        "Continuous-reasoning language models replace long written chains of thought with a small number of continuous hidden states. "
        "This reduces the number of generated reasoning tokens, but the model must still convert its final 768-dimensional state into scores for all 50,257 vocabulary tokens. "
        "Low-rank vocabulary heads are an established acceleration technique: most directly, SlimSpec replaces a speculative drafter's full head with two learned projections, preserves the complete vocabulary, and reports approximately 4-5-fold head acceleration with smaller end-to-end gains [10]. "
        "This proposal addresses a different question: whether the answer-bearing part of a frozen CODI model's state can be identified through covariance-matrix eigendecomposition and causal intervention, then used to construct a post-training, lower-cost vocabulary readout. "
        "Preliminary experiments on the released CODI GPT-2 checkpoint identified 28 eigendirections, indexed 4 through 31, at the final answer-cue state. "
        "Although these directions explain only 11.31 percent of state variance, retaining them preserves 87.8 percent of baseline numeric exact-match accuracy, while removing them reduces accuracy by 30.48 percentage points. "
        "A rank-32 variant preserves 94.4 percent of baseline accuracy. By projecting the vocabulary matrix into the selected eigenspace offline, the online readout can operate on 28 or 32 coordinates rather than 768. "
        "An isolated Tesla T4 benchmark measured the full head at 1,327 microseconds and the rank-32 head at 117 microseconds, an 11.3-fold component-level speedup. "
        "The proposed research will integrate this readout into end-to-end generation, compare it directly with learned SlimSpec-style and post-training compression baselines, and test whether behavior-selected eigendirections transfer across answer positions, tasks, and checkpoints."
    ), after=9)
    abstract.paragraph_format.left_indent = Inches(0.5)
    abstract.paragraph_format.right_indent = Inches(0.5)

    add_heading(doc, "1. Introduction", 1)
    add_para(doc, (
        "Language models normally express reasoning as text. Each reasoning token requires another pass through the transformer and another vocabulary prediction. CODI, or Continuous Chain-of-Thought via Self-Distillation, instead learns a short sequence of continuous thoughts and aligns the student's internal answer representation with that of an explicit-reasoning teacher [1]. "
        "The released CODI system used in this project is based on GPT-2 small [2]. It contains twelve transformer blocks, has a representation width of 768, and uses a vocabulary of 50,257 tokens. During CODI inference, a question is encoded, six continuous latent iterations are executed through the shared transformer, and the model then receives the answer cue. The final representation at the colon in 'The answer is:' is transformed by the language-model head into one score for every vocabulary token."
    ))
    add_para(doc, (
        "The language-model head is therefore a natural optimization target. Its learned vocabulary matrix can be viewed as 50,257 token-specific detectors, each acting on the 768 coordinates of the final hidden state. A normal readout calculates a dense matrix product for every generated token. Prior work already shows that a learned low-rank factorization can accelerate a speculative draft head [10]. The open question here is narrower and mechanistic: can a frozen continuous-reasoning model be compressed after training by selecting directions according to their causal contribution to answer accuracy rather than learning an unconstrained low-rank bottleneck?"
    ))

    add_heading(doc, "2. Problem statement", 1)
    add_para(doc, (
        "A low-rank language-model head is not itself a new contribution. SlimSpec learns a factorized full-vocabulary head for an EAGLE-3 speculative drafter and demonstrates that head latency can be reduced substantially when acceptance quality is preserved [10]. However, its factors are learned during drafter training, its rank is selected manually, and a full target model verifies the proposed tokens. It does not identify which directions in a frozen model are causally responsible for task accuracy."
    ))
    add_para(doc, (
        "Our preliminary CODI result exposes a gap between variance and answer contribution. The first four eigendirections contain 82.31 percent of variance but preserve only 6.1 percent of baseline exact-match accuracy when retained alone. Eigdirections 4 through 31 contain only 11.31 percent of variance but preserve 87.8 percent of baseline accuracy. The research problem is therefore to determine whether behavior-selected directions can support a practical post-training readout for the main answer-generating model. Unlike speculative decoding, CODI has no unchanged target model that can correct a bad proposal, so accuracy loss must be measured directly and kept within a declared tolerance."
    ))

    add_heading(doc, "3. Aim, objectives, and research questions", 1)
    add_para(doc, "Aim", size=10, bold=True, keep=True, after=3)
    add_para(doc, (
        "To develop and evaluate a post-training, behavior-selected eigenspace method for reducing CODI's vocabulary-readout cost, and to establish when its accuracy-latency trade-off is preferable to a learned low-rank head."
    ))
    add_para(doc, "Objectives", size=10, bold=True, keep=True, after=3)
    add_number(doc, "Identify stable answer-bearing eigendirections at CODI's final answer-cue representation using held-out calibration and causal retention/removal interventions.")
    add_number(doc, "Convert the selected eigenspace into rank-28 and rank-32 language-model heads through an offline transformation of the learned vocabulary matrix.")
    add_number(doc, "Measure exact-match accuracy, token-ranking agreement, numerical error, memory use, head-cost fraction, and end-to-end latency.")
    add_number(doc, "Compare behavior-selected eigendecomposition with a learned SlimSpec-style factorization, weight-only decomposition, variance-ranked selection, activation-aware compression, and vocabulary shortlisting.")
    add_number(doc, "Test whether one endpoint-derived eigenspace remains valid across later answer positions, data splits, random seeds, mathematical reasoning datasets, and compatible checkpoints.")
    add_para(doc, "Research questions", size=10, bold=True, keep=True, before=6, after=3)
    add_bullet(doc, "How many eigendirections are sufficient and necessary for CODI's answer accuracy at the forced answer cue?")
    add_bullet(doc, "Does behavior-based eigendirection selection preserve accuracy better than selection based only on variance or matrix reconstruction error?")
    add_bullet(doc, "At equal rank and operation count, how does the fixed eigenspace head compare with a learned full-vocabulary low-rank head?")
    add_bullet(doc, "Does the reduction in language-model-head operations produce a measurable end-to-end inference improvement?")
    add_bullet(doc, "Are the selected eigendirections stable across generated-token positions, calibration samples, tasks, and model checkpoints?")
    add_bullet(doc, "Does the behavior-selected eigenspace method generalize beyond CODI to other continuous- or latent-reasoning models and to conventional autoregressive language models with different architectures, representation widths, and vocabularies?")

    add_heading(doc, "4. Background and related work", 1)
    add_heading(doc, "4.1 Continuous reasoning and CODI", 2)
    add_para(doc, (
        "CODI trains a shared model in teacher and student modes. The teacher receives an explicit chain of thought, whereas the student performs a fixed number of continuous latent iterations. Training combines language-model objectives with alignment between teacher and student representations at the final answer-generating position [1]. At inference, the teacher branch is not required: the student processes the question, performs the latent iterations, consumes the answer cue, and generates answer tokens through the normal vocabulary head."
    ))
    add_heading(doc, "4.2 Low-rank and activation-aware compression", 2)
    add_para(doc, (
        "Low-rank factorization replaces one large linear map with two smaller maps. Fisher-weighted factorization showed that minimizing ordinary reconstruction error can be misaligned with task accuracy and proposed weighting parameters by their effect on model predictions [4]. SliceGPT uses orthogonal transformations and post-training dimensional reduction across transformer representations [5]. ASVD adapts matrix decomposition to observed activation distributions [6], while FLAT-LLM performs training-free transformations using eigendirections of activation spaces [7]. These studies establish that model behavior and activation statistics should influence compression, but they do not test CODI's answer-cue endpoint or the particular non-leading eigenspace identified here."
    ))
    add_heading(doc, "4.3 SlimSpec and vocabulary-head acceleration", 2)
    add_para(doc, (
        "SlimSpec is the closest systems precedent [10]. For a draft state h of width d and vocabulary size V, it replaces z = W_full h with two learned maps:"
    ))
    add_formula(doc, "z = W_up W_down h")
    add_para(doc, (
        "Here W_down has shape r by d and W_up has shape V by r, reducing complexity from O(Vd) to O(rd + Vr) while retaining all vocabulary tokens. With EAGLE-3 drafters for Llama-3.1-8B, GPT-OSS-20B, and Qwen3-30B-A3B, SlimSpec reports approximately 4-5-fold draft-head acceleration. Its default rank is d/8, and representative Llama results retain about 99 percent of average draft-token acceptance. End-to-end improvements are smaller and depend on how much of the full speculative pipeline is occupied by the head."
    ))
    add_para(doc, (
        "This evidence establishes that dense full-vocabulary factorization is a credible engineering strategy, but the setting differs from the present study. SlimSpec trains the two factors as part of a speculative drafter; an unchanged target model subsequently verifies its tokens. The proposed CODI head is derived after training from a frozen main model, and its factors are constrained by an activation eigenspace selected through answer-accuracy interventions. Because there is no verifier, any altered token ranking can directly alter final answer accuracy."
    ))
    add_heading(doc, "4.4 Other output-head methods", 2)
    add_para(doc, (
        "Adaptive softmax reduces expected vocabulary computation through token-frequency structure [8], and vocabulary clustering predicts a smaller candidate set before projection [9]. ARCHead combines a compressed output head with an activation-derived metric [11]. These methods, together with SlimSpec, mean that the proposed work should not claim invention of low-rank output projection. Its intended contribution is the causal selection, post-training construction, and end-to-end evaluation of a non-leading answer-bearing eigenspace in CODI."
    ))

    add_heading(doc, "5. Preliminary work completed", 1)
    add_heading(doc, "5.1 Reproduction and state collection", 2)
    add_para(doc, (
        "The author-released zen-E/CODI-gpt2 checkpoint was reproduced in a pinned software environment. Native GSM8K numeric exact match was 0.4359. Under the forced answer cue used for intervention, baseline accuracy was 0.4337, corresponding to 572 correct answers among 1,319 test questions. All confirmatory runs used float32. For eigenspace construction, final answer-cue representations were collected from 2,048 GSM8K training questions. Each representation was the 768-dimensional state labelled state 12, produced after the twelve GPT-2 blocks at the colon in the answer cue and immediately before vocabulary readout."
    ))
    add_heading(doc, "5.2 Covariance eigendecomposition", 2)
    add_para(doc, (
        "Let each collected endpoint representation be h_i and let mu be their mean. The centered rows were stacked into a matrix X. The 768 by 768 covariance matrix was then constructed as:"
    ))
    add_formula(doc, "C = (1 / (n - 1)) X^T X")
    add_para(doc, "Because C is symmetric, it can be decomposed into orthonormal eigendirections and their associated eigenvalues:")
    add_formula(doc, "C = U Lambda U^T")
    add_para(doc, (
        "The columns of U define directions in the original 768-dimensional representation space. The eigenvalues in Lambda measure how strongly the calibration states vary along those directions. Retention interventions replaced the unselected component of a test state with the calibration mean. Removal interventions did the reverse. Full greedy decoding was then used to measure numeric exact match, so the result reflects complete answers rather than only the first generated token."
    ))
    add_heading(doc, "5.3 Discovery of 28 answer-bearing directions", 2)
    add_results_table(doc)
    add_para(doc, (
        "All preregistered confirmation conditions passed. Retaining eigendirections 4-31 produced 0.3806 exact-match accuracy, equal to 87.8 percent of the 0.4337 baseline. Removing the same 28 directions reduced accuracy to 0.1289, a loss of 30.48 percentage points; the paired confidence interval was +27.90 to +33.13 points and the exact McNemar p-value was 5.7 x 10^-101. Four matched random rank-28 subspaces retained only 6.5 to 8.7 percent of baseline accuracy. The selected band was also stable across disjoint calibration halves, with mean principal-angle cosines of 0.979 and 0.968."
    ))
    add_para(doc, (
        "These results establish a bounded mechanistic claim: at this frozen CODI checkpoint, dataset, endpoint, and intervention type, a 28-dimensional linear subspace is both highly sufficient and strongly necessary for answer accuracy. The experiment does not establish that the same directions are universal, nonlinear, or independently interpretable. Retention also moves states away from their ordinary distribution because discarded coordinates are replaced by the mean."
    ))

    add_heading(doc, "6. Proposed eigenspace readout", 1)
    add_heading(doc, "6.1 Standard language-model head", 2)
    add_para(doc, (
        "Let h be the final 1 by 768 hidden state and W be the learned 50,257 by 768 vocabulary matrix. The standard head calculates:"
    ))
    add_formula(doc, "z = h W^T")
    add_para(doc, (
        "The result z contains 50,257 logits. Each logit is a score for one vocabulary token. Greedy decoding selects the token with the largest score. This multiplication requires 768 x 50,257 = 38,597,376 multiply-accumulate operations for one state."
    ))
    add_heading(doc, "6.2 Offline head transformation", 2)
    add_para(doc, (
        "Choose r columns of U and place them in U_r, where r is 28 or 32. Before inference, project every vocabulary detector into the selected eigenspace:"
    ))
    add_formula(doc, "A = W U_r")
    add_para(doc, (
        "For rank 28, A has shape 50,257 by 28. For rank 32, it has shape 50,257 by 32. The centering contribution is also computed once as b = mu W^T. Neither A nor b depends on the incoming question, so both can be stored with the optimized model."
    ))
    add_heading(doc, "6.3 Online inference", 2)
    add_para(doc, "For each final hidden state, inference performs two smaller operations:")
    add_formula(doc, "c = (h - mu) U_r")
    add_formula(doc, "z_hat = b + c A^T")
    add_para(doc, (
        "The first operation converts 768 coordinates into r eigenspace coordinates. The second converts those coordinates into scores for the full vocabulary. The full vocabulary remains available; the approximation restricts how hidden-state information reaches the vocabulary matrix. Rank 28 corresponds exactly to the discovered answer-bearing band 4-31. Rank 32 additionally includes directions 0-3 and achieved higher retained accuracy in the preliminary intervention."
    ))
    add_efficiency_table(doc)
    add_para(doc, (
        "The theoretical multiply-accumulate reduction does not translate one-for-one into elapsed-time reduction because memory movement, kernel launch cost, batching, and the transformer blocks also consume time. Nevertheless, the component benchmark shows that the factorized computation can be materially faster when implemented as dense matrix operations."
    ))
    add_heading(doc, "6.4 Relationship to SlimSpec", 2)
    add_para(doc, (
        "The proposed readout and SlimSpec share the same two-stage computational skeleton but obtain their factors differently. Under column-vector notation, SlimSpec learns W_down and W_up freely. The fixed CODI construction is the constrained special case W_down = U_r^T and W_up = W U_r, with an additional centering bias. Thus the comparison tests whether causal behavioral selection supplies a useful inductive bias or whether unconstrained low-rank training is superior. SlimSpec commonly evaluates r = d/8; for CODI, d/8 is rank 96, whereas ranks 28 and 32 are much more aggressive bottlenecks, approximately d/27 and d/24."
    ))
    add_heading(doc, "6.5 End-to-end cost model", 2)
    add_para(doc, (
        "Following SlimSpec's separation of head and non-head time [10], let nu be the optimized-to-full head latency ratio and let kappa be full-head time divided by all non-head time. If output length is unchanged, the ideal end-to-end speedup attributable to the replacement is:"
    ))
    add_formula(doc, "S = (1 + kappa) / (1 + nu kappa)")
    add_para(doc, (
        "SlimSpec additionally multiplies this term by draft-token acceptance preservation because acceptance controls useful tokens per speculative round. Plain CODI decoding has no verifier and no corresponding acceptance mechanism; therefore exact-match accuracy and token-ranking fidelity must remain separate quality constraints rather than being folded into the speed formula. Measured end-to-end latency remains authoritative because the approximation may change answer length or kernel behavior."
    ))
    add_heading(doc, "6.6 Latent-budget result", 2)
    add_para(doc, (
        "A separate preliminary efficiency experiment reduced CODI's six continuous latent iterations. On full GSM8K, the trained M=6 setting scored 0.4337, while M=5 scored 0.4359, a difference of +0.23 percentage points. M=4 and M=3 fell to 0.3836 and 0.3768. Thus one latent iteration could be removed without an observed accuracy cost, but batch-32 wall-clock time on the T4 was essentially flat because fixed costs dominated. This result will be re-evaluated under latency-bound batch-1 inference and combined with the low-rank readout only after each intervention is validated independently."
    ))

    add_heading(doc, "7. Proposed methodology", 1)
    add_heading(doc, "7.1 Experimental design", 2)
    add_bullet(doc, "Freeze independent training/calibration, eigenspace-selection, validation, and final test partitions to prevent using test performance to choose directions or rank.")
    add_bullet(doc, "Collect answer-cue state-12 representations and compute the covariance matrix and its eigendecomposition in float32.")
    add_bullet(doc, "Select candidate direction bands using only the selection partition and preregister the final rank and intervention thresholds.")
    add_bullet(doc, "Construct rank-28, rank-32, and rank-sweep heads offline. Evaluate endpoint-only replacement, one shared head across all answer positions, and position-specific or fallback variants before selecting a deployment rule.")
    add_bullet(doc, "Integrate the selected head into the released CODI generation loop and separately time transformer, latent-iteration, full-head, compressed-head, sampling, and framework overhead components.")
    add_bullet(doc, "Evaluate correctness and efficiency on untouched test sets. Repeat the process across seeds, compatible checkpoints, and additional mathematical reasoning datasets.")

    add_heading(doc, "7.2 Comparison methods", 2)
    add_bullet(doc, "The original full 768 by 50,257 vocabulary head.")
    add_bullet(doc, "Leading-eigenvalue rank-r selection without behavioral intervention.")
    add_bullet(doc, "Weight-only truncated decomposition of the vocabulary matrix.")
    add_bullet(doc, "Activation-aware and Fisher-weighted low-rank baselines.")
    add_bullet(doc, "A learned SlimSpec-style full-vocabulary factorization trained by logit distillation at matched ranks 28, 32, 64, and 96, with both random and eigenspace initialization.")
    add_bullet(doc, "Vocabulary shortlisting or clustering as a different approach to the same output bottleneck.")

    add_heading(doc, "7.3 Outcome measures", 2)
    add_bullet(doc, "Primary quality outcome: numeric exact-match accuracy with paired confidence intervals.")
    add_bullet(doc, "Readout fidelity: top-1 and top-k token agreement, logit error, cross-entropy, answer-margin preservation, and agreement by generated-token position.")
    add_bullet(doc, "Efficiency: batch-1 and batch-32 latency, tokens per second, peak accelerator memory, head time, complete generation time, nu, kappa, and predicted-versus-measured speedup.")
    add_bullet(doc, "Robustness: eigenspace overlap, retained accuracy, and latency across calibration resamples, tasks, seeds, and checkpoints.")

    add_heading(doc, "7.4 Statistical analysis", 2)
    add_para(doc, (
        "All accuracy comparisons will be paired by question. Exact McNemar tests will measure discordant correctness outcomes, and paired bootstrap intervals will quantify accuracy differences and retention ratios. Multiple planned comparisons will use a declared correction procedure. Direction selection will not be changed after the test set is opened. Hardware timing will include warm-up, synchronization, repeated trials, medians, dispersion, and fixed software and precision settings."
    ))

    add_heading(doc, "8. Expected contributions", 1)
    add_bullet(doc, "A mechanistic account of where CODI's final answer information is concentrated at its answer cue.")
    add_bullet(doc, "A behavior-selected alternative to compression based only on eigenvalue magnitude or reconstruction error.")
    add_bullet(doc, "A post-training, constrained full-vocabulary factorization derived from a frozen continuous-reasoning model rather than learned as a speculative draft head.")
    add_bullet(doc, "A controlled comparison with SlimSpec-style learned factorization at equal rank and computational cost.")
    add_bullet(doc, "A careful separation among causal sufficiency, token-level fidelity, component acceleration, and real end-to-end inference improvement.")
    add_bullet(doc, "A reproducible evaluation protocol for connecting hidden-state interventions to systems optimization.")

    add_heading(doc, "9. Risks and limitations", 1)
    add_para(doc, (
        "The strongest preliminary evidence is limited to one released CODI GPT-2 checkpoint, GSM8K, one semantic endpoint, and linear interventions. A subspace selected on one task can overfit task vocabulary or answer format, and mean replacement can create hidden states that are not naturally produced by the model. The rank-28 readout loses approximately 5.3 absolute accuracy points relative to baseline, while rank 32 loses approximately 2.4 points. Unlike SlimSpec's speculative drafter, the CODI head has no unchanged target model to reject incorrect proposals, so compression errors directly affect answers. Evidence collected at the colon may also fail at later generated positions. The isolated 11.3-fold T4 head speedup cannot be compared directly with SlimSpec's H200 results or reported as an end-to-end CODI speedup. Finally, low-rank LM heads are established prior art; novelty depends on causal direction selection, frozen post-training construction, and convincing matched-baseline evaluation."
    ))

    add_heading(doc, "10. Conclusion", 1)
    add_para(doc, (
        "SlimSpec establishes that a learned two-stage, full-vocabulary head can accelerate speculative drafting, so the contribution proposed here is not low-rank projection by itself. The preliminary CODI study instead found that a non-leading set of 28 covariance eigendirections carries a small share of total variance but most of the frozen model's answer accuracy. This supports a distinct hypothesis: causal behavioral selection may construct a useful post-training factorization of a continuous-reasoning model's main output head. The decisive tests are comparison with a learned SlimSpec-style head at matched rank, validation across every answer position, and complete latency measurement under an explicit accuracy constraint. The proposed research will determine whether the mechanistic observation becomes a defensible inference method rather than an isolated component benchmark."
    ))

    doc.add_page_break()
    add_heading(doc, "References", 1)
    refs = [
        ("Z. Shen, H. Yan, L. Zhang, Z. Hu, Y. Du, and Y. He, 'CODI: Compressing Chain-of-Thought into Continuous Space via Self-Distillation,' arXiv:2502.21074, 2025.", "https://arxiv.org/abs/2502.21074"),
        ("A. Radford et al., 'Language Models are Unsupervised Multitask Learners,' OpenAI technical report, 2019.", "https://cdn.openai.com/better-language-models/language_models_are_unsupervised_multitask_learners.pdf"),
        ("K. Cobbe et al., 'Training Verifiers to Solve Math Word Problems,' arXiv:2110.14168, 2021.", "https://arxiv.org/abs/2110.14168"),
        ("Y.-C. Hsu, T. Hua, S. Chang, Q. Lou, Y. Shen, and H. Jin, 'Language Model Compression with Weighted Low-Rank Factorization,' ICLR, 2022.", "https://arxiv.org/abs/2207.00112"),
        ("S. Ashkboos, M. L. Croci, M. Gennari do Nascimento, T. Hoefler, and J. Hensman, 'SliceGPT: Compress Large Language Models by Deleting Rows and Columns,' ICLR, 2024.", "https://proceedings.iclr.cc/paper_files/paper/2024/hash/316648eb8b4ffb6010f531b07848c300-Abstract-Conference.html"),
        ("Z. Yuan, Y. Shang, Y. Song, D. Yang, Q. Wu, Y. Yan, and G. Sun, 'ASVD: Activation-aware Singular Value Decomposition for Compressing Large Language Models,' arXiv:2312.05821, 2023, revised 2025.", "https://arxiv.org/abs/2312.05821"),
        ("J. Tian, R. Solgi, J. Lu, Y. Yang, H. Li, and Z. Zhang, 'FLAT-LLM: Fine-grained Low-rank Activation Space Transformation for Large Language Model Compression,' arXiv:2505.23966, 2025, revised 2026.", "https://arxiv.org/abs/2505.23966"),
        ("E. Grave, A. Joulin, M. Cisse, D. Grangier, and H. Jegou, 'Efficient Softmax Approximation for GPUs,' Proceedings of ICML, pp. 1302-1310, 2017.", "https://proceedings.mlr.press/v70/grave17a.html"),
        ("H. Amer, M. Afify, Y. J. Kim, H. Matsushita, and H. Hassan, 'Fast Vocabulary Projection Method via Clustering for Multilingual Machine Translation on GPU,' Proceedings of AMTA, pp. 58-69, 2022.", "https://aclanthology.org/2022.amta-research.5/"),
        ("A. Plaksin, S. Krutikov, S. Skvortsov, and A. Samarin, 'SlimSpec: Low-Rank Draft LM-Head for Accelerated Speculative Decoding,' arXiv:2605.10453, 2026.", "https://arxiv.org/abs/2605.10453"),
        ("S. T. Kocabay, T. R. Akkus, and K. A. Yuksel, 'ARCHead: Activation-Metric Residual Correction for Large Language Model Output Heads,' arXiv:2608.02703, 2026.", "https://arxiv.org/abs/2608.02703"),
        ("M. J. Raza, 'Official CODI Endpoint Eigenspace Confirmation and Efficiency Measurements,' unpublished project experiment records, 2026.", "Repository experiment ledger and reproducibility artifacts."),
    ]
    for i, (citation, url) in enumerate(refs, 1):
        add_reference(doc, i, citation, url)

    # NeurIPS-style non-anonymous preprint notice on the first page.
    footer = section.first_page_footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Research proposal. Work in progress."), size=8)

    # Enforce black font for every run, including table text and footer fields.
    for paragraph in list(doc.paragraphs):
        for run in paragraph.runs:
            run.font.color.rgb = BLACK
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = BLACK

    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build()
